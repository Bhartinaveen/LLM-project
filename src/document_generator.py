"""
Document Generator Module
Converts LLM-generated content to formatted DOCX files
"""

import logging
import re
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate DOCX documents from LLM content"""

    def __init__(self, output_dir: str = "./outputs"):
        """
        Initialize document generator
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"DocumentGenerator initialized with output dir: {output_dir}")

    def generate_document(
        self, content: str, document_type: str, metadata: Optional[dict] = None
    ) -> str:
        """
        Generate DOCX document from LLM content
        
        Args:
            content: LLM-generated document content
            document_type: Type of document (loan_agreement, etc.)
            metadata: Optional metadata dict with document info
            
        Returns:
            Path to generated document
        """
        doc = Document()

        # Add content sections
        self._add_document_content(doc, content, metadata)

        # Generate filename
        filename = self._generate_filename(document_type)
        filepath = self.output_dir / filename

        # Save document
        doc.save(str(filepath))
        logger.info(f"Document saved: {filepath}")

        return str(filepath)

    def _add_document_content(
        self, doc: Document, content: str, metadata: Optional[dict] = None
    ) -> None:
        """
        Add content to document with formatting
        
        Args:
            doc: Document object
            content: Document content
            metadata: Optional metadata
        """
        lines = content.split("\n")
        
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Heading 1 (# Title)
            if line.startswith("# "):
                self._add_heading(doc, line[2:].strip(), level=1)
            
            # Heading 2 (## Section)
            elif line.startswith("## "):
                self._add_heading(doc, line[3:].strip(), level=2)
            
            # Heading 3 (### Subsection)
            elif line.startswith("### "):
                self._add_heading(doc, line[4:].strip(), level=3)
            
            # Signature Block Placeholder
            elif "[SIGNATURE_BLOCK]" in line:
                self._add_signature_block(doc)
            
            # List items (Bullet points)
            elif line.startswith("- ") or line.startswith("* "):
                self._add_paragraph(doc, line[2:].strip(), style="List Bullet")
            
            # Numbered lists (1. Item)
            elif re.match(r"^\d+\.\s+", line):
                match = re.match(r"^\d+\.\s+(.+)$", line)
                if match:
                    self._add_paragraph(doc, match.group(1).strip(), style="List Number")
            
            # Regular Paragraphs
            else:
                self._add_paragraph(doc, line)

        # Add metadata footer
        if metadata:
            self._add_footer(doc, metadata)

    def _add_heading(self, doc: Document, heading: str, level: int = 1) -> None:
        """Add heading to document with style"""
        if level == 1:
            # Main Title Style
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "Arial"
        else:
            # Section Headings
            p = doc.add_heading(heading, level=level)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_paragraph(self, doc: Document, text: str, style: str = None) -> None:
        """Add paragraph to document with bold support"""
        if style:
            p = doc.add_paragraph(style=style)
        else:
            p = doc.add_paragraph()
        
        # Parse bold markdown (**text**)
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(part)
            
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

    def _add_signature_block(self, doc: Document) -> None:
        """Add professional signature block"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        for cell in table.rows[0].cells:
            cell.width = Inches(3.0)

        # Lender/Landlord Signature
        c1 = table.rows[0].cells[0].paragraphs[0]
        c1.add_run("__________________________\n").font.bold = True
        c1.add_run("Signed by (Party A)\n")
        c1.add_run("Date: _____________")

        # Borrower/Tenant Signature
        c2 = table.rows[0].cells[1].paragraphs[0]
        c2.add_run("__________________________\n").font.bold = True
        c2.add_run("Signed by (Party B)\n")
        c2.add_run("Date: _____________")
        
        doc.add_paragraph()

    def _add_footer(self, doc: Document, metadata: dict) -> None:
        """Add footer with metadata"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.style.font.size = Pt(8)

        # Add metadata as footer text
        if "document_type" in metadata:
            footer_para.text += f" | Type: {metadata['document_type']}"

    def _generate_filename(self, document_type: str) -> str:
        """
        Generate filename for document
        
        Args:
            document_type: Type of document
            
        Returns:
            Filename with timestamp
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_type_name = (
            document_type.replace("_", "-").title()
            if "_" in document_type
            else document_type
        )
        return f"{doc_type_name}_{timestamp}.docx"

    def add_cover_page(
        self, doc: Document, title: str, parties: list, date: str
    ) -> None:
        """
        Add professional cover page
        
        Args:
            doc: Document object
            title: Document title
            parties: List of parties involved
            date: Document date
        """
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(20)
            run.font.bold = True

        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Parties
        parties_para = doc.add_paragraph("Between:")
        for run in parties_para.runs:
            run.font.bold = True

        for party in parties:
            doc.add_paragraph(party, style="List Bullet")

        # Date
        doc.add_paragraph()
        date_para = doc.add_paragraph(f"Date: {date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        doc.add_page_break()


def generate_legal_document(
    content: str,
    document_type: str,
    output_dir: str = "./outputs",
    metadata: Optional[dict] = None,
) -> str:
    """
    Convenience function to generate legal document
    
    Args:
        content: LLM-generated content
        document_type: Type of document
        output_dir: Output directory
        metadata: Optional metadata
        
    Returns:
        Path to generated DOCX file
    """
    generator = DocumentGenerator(output_dir)
    return generator.generate_document(content, document_type, metadata)
